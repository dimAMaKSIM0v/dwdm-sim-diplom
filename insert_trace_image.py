# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches
from pathlib import Path

src = Path(r"C:\Users\dimam\Desktop\NIR_updated_v4.docx")
out = Path(r"C:\Users\dimam\Desktop\NIR_updated_v5.docx")
img = Path(r"C:\Users\dimam\Desktop\111.png")

if not img.exists():
    raise SystemExit(f"image not found: {img}")

# Helper to insert image after a paragraph
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    if style:
        p.style = style
    return p


doc = Document(src)

# find 2.3 section and the paragraph about OSRM / straight line
anchor = None
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("Трассировка линии выполняется либо по дорожной сети"):
        anchor = p
        break

if anchor is None:
    raise SystemExit("Anchor paragraph not found")

# Insert image paragraph after anchor
p_img = insert_paragraph_after(anchor)
run = p_img.add_run()
run.add_picture(str(img), width=Inches(5.6))

# Insert caption after image
p_cap = insert_paragraph_after(p_img, "Рисунок 5 — Пример трассировки линий на карте топологии")

# Save
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(str(out))
